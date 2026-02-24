"""组卷 Word 导出"""
from __future__ import annotations
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from med_exam_toolkit.models import Question
from med_exam_toolkit.exam.config import ExamConfig

# AI 补全答案/解析的提示颜色
_AI_COLOR = RGBColor(0xCC, 0x77, 0x00)   # 橙色，区别于正式答案的绿色

class ExamDocxExporter:

    def __init__(self, config: ExamConfig):
        self.config = config

    def export(self, questions: list[Question], output_path: Path) -> Path:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        fp = output_path.with_suffix(".docx")

        doc = Document()
        self._set_default_font(doc)
        self._add_header(doc, questions)
        self._add_questions(doc, questions)

        if self.config.answer_sheet:
            doc.add_page_break()
            self._add_answer_sheet(doc, questions)

        doc.save(str(fp))
        print(f"[INFO] 试卷导出完成: {fp} ({len(questions)} 题)")
        return fp

    # ── 页面设置 ──

    @staticmethod
    def _set_font(run, name: str = "宋体", size: Pt | None = None):
        """同时设置中西文字体"""
        run.font.name = name
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), name)
        if size is not None:
            run.font.size = size

    @staticmethod
    def _set_default_font(doc: Document):
        style = doc.styles["Normal"]
        style.font.name = "宋体"
        style.font.size = Pt(10.5)
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.paragraph_format.space_after = Pt(2)
        style.paragraph_format.line_spacing = 1.25
        for section in doc.sections:
            section.top_margin    = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin   = Cm(2.8)
            section.right_margin  = Cm(2.8)

    # ── 试卷头 ──

    def _add_header(self, doc: Document, questions: list[Question]):
        cfg = self.config
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(cfg.title)
        run.bold = True
        run.font.size = Pt(18)

        if cfg.subtitle:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(cfg.subtitle)
            run.font.size = Pt(12)

        info_parts = []
        if cfg.time_limit:
            info_parts.append(f"考试时间: {cfg.time_limit} 分钟")
        if cfg.total_score:
            info_parts.append(f"满分: {cfg.total_score} 分")
        info_parts.append(f"共 {len(questions)} 题")
        if cfg.score_per_sub:
            info_parts.append(f"每题 {cfg.score_per_sub} 分")

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("    ".join(info_parts))
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        doc.add_paragraph("━" * 43)

    # ── 题目区 ──

    def _add_questions(self, doc: Document, questions: list[Question]):
        current_mode = ""
        global_idx   = 0

        for q in questions:
            if q.mode != current_mode:
                current_mode = q.mode
                p = doc.add_paragraph()
                p.space_before = Pt(12)
                run = p.add_run(f"【{current_mode}】")
                run.bold = True
                run.font.size = Pt(12)

            if len(q.sub_questions) > 1 or q.stem:
                global_idx = self._add_compound_question(doc, q, global_idx)
            else:
                global_idx += 1
                self._add_single_question(doc, q.sub_questions[0], global_idx)

    def _add_single_question(self, doc: Document, sq, idx: int):
        p = doc.add_paragraph()
        run = p.add_run(f"{idx}. {sq.text}")
        run.font.size = Pt(10.5)

        self._add_options(doc, sq.options)

        if self.config.show_answers:
            self._add_inline_answer(doc, sq)

        doc.add_paragraph("")

    def _add_compound_question(self, doc: Document, q: Question, global_idx: int) -> int:
        stem = q.stem or q.raw.get("test", "")
        if stem:
            p = doc.add_paragraph()
            p.space_before = Pt(6)
            run = p.add_run(
                f"（{global_idx + 1}～{global_idx + len(q.sub_questions)} 题共用题干）"
            )
            run.bold = True
            run.font.size = Pt(10)

            p = doc.add_paragraph()
            run = p.add_run(stem)
            run.font.size = Pt(10.5)
            run.italic = True

        if q.shared_options:
            self._add_options(doc, q.shared_options)

        for sq in q.sub_questions:
            global_idx += 1
            p = doc.add_paragraph()
            run = p.add_run(f"{global_idx}. {sq.text}")
            run.font.size = Pt(10.5)

            if not q.shared_options:
                self._add_options(doc, sq.options)

            if self.config.show_answers:
                self._add_inline_answer(doc, sq)

        doc.add_paragraph("")
        return global_idx

    def _add_inline_answer(self, doc: Document, sq) -> None:
        """题目正文中内嵌答案（show_answers 模式）"""
        ans = sq.eff_answer
        if not ans:
            return
        p = doc.add_paragraph()
        run = p.add_run(f"【答案】{ans}")
        run.font.size = Pt(9)
        if sq.answer_source == "ai":
            run.font.color.rgb = _AI_COLOR
            run.add_run("(AI)")  # 不加括号，直接标注
        else:
            run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)

    @staticmethod
    def _display_width(text: str) -> int:
        """计算显示宽度：CJK 字符算 2，其余算 1"""
        w = 0
        for ch in text:
            if '\u4e00' <= ch <= '\u9fff' or '\u3000' <= ch <= '\u303f' or '\uff00' <= ch <= '\uffef':
                w += 2
            else:
                w += 1
        return w

    @staticmethod
    def _add_options(doc: Document, options: list[str]):
        if not options:
            return

        max_display = max(ExamDocxExporter._display_width(o) for o in options)

        if max_display <= 28 and len(options) <= 6:
            tab_pos = Cm(9)
            for i in range(0, len(options), 2):
                left  = options[i]
                right = options[i + 1] if i + 1 < len(options) else ""
                p = doc.add_paragraph()
                p.paragraph_format.space_after  = Pt(0)
                p.paragraph_format.space_before = Pt(0)

                pPr  = p._element.get_or_add_pPr()
                tabs = OxmlElement("w:tabs")
                tab  = OxmlElement("w:tab")
                tab.set(qn("w:val"), "left")
                tab.set(qn("w:pos"), str(int(tab_pos.emu / 635)))
                tabs.append(tab)
                pPr.append(tabs)

                run = p.add_run(f"    {left}")
                run.font.size = Pt(10)
                if right:
                    run = p.add_run("\t")
                    run.font.size = Pt(10)
                    run = p.add_run(right)
                    run.font.size = Pt(10)
        else:
            for opt in options:
                p = doc.add_paragraph()
                run = p.add_run(f"    {opt}")
                run.font.size = Pt(10)
                p.paragraph_format.space_after  = Pt(0)
                p.paragraph_format.space_before = Pt(0)

    # ── 答案页 ──
    def _add_answer_sheet(self, doc: Document, questions: list[Question]):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("参考答案")
        run.bold = True
        run.font.size = Pt(14)
        doc.add_paragraph("")

        if self.config.show_discuss:
            current_mode = ""
            idx = 0
            for q in questions:
                if q.mode != current_mode:
                    current_mode = q.mode
                    p = doc.add_paragraph()
                    p.space_before = Pt(8)
                    run = p.add_run(f"【{current_mode}】")
                    run.bold = True
                    run.font.size = Pt(11)

                for sq in q.sub_questions:
                    idx += 1
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after  = Pt(1)
                    p.paragraph_format.space_before = Pt(1)

                    run = p.add_run(f"{idx}. ")
                    run.font.size = Pt(10)

                    # 答案（AI 兜底时橙色标注）
                    ans = sq.eff_answer
                    run = p.add_run(ans if ans else "—")
                    run.bold = True
                    run.font.size = Pt(10)
                    if sq.answer_source == "ai":
                        run.font.color.rgb = _AI_COLOR
                    else:
                        run.font.color.rgb = RGBColor(0x00, 0x66, 0x00)

                    if sq.answer_source == "ai":
                        run = p.add_run("(AI)")
                        run.font.size = Pt(8)
                        run.font.color.rgb = _AI_COLOR

                    # 解析（AI 兜底时橙色标注）
                    dis = sq.eff_discuss
                    if dis:
                        run = p.add_run(f"  {dis}")
                        run.font.size = Pt(9)
                        if sq.discuss_source == "ai":
                            run.font.color.rgb = _AI_COLOR
                        else:
                            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

            doc.add_paragraph("")

        self._add_quick_ref(doc, questions)

    def _add_quick_ref(self, doc: Document, questions: list[Question]):
        """答案速查表 — 等宽对齐，AI 答案加 * 标注"""
        p = doc.add_paragraph()
        run = p.add_run("答案速查")
        run.bold = True
        run.font.size = Pt(11)

        answers = []
        idx = 0
        for q in questions:
            for sq in q.sub_questions:
                idx += 1
                ans = sq.eff_answer or "—"
                # AI 来源在速查表里用 * 后缀标注
                if sq.answer_source == "ai":
                    ans = ans + "*"
                answers.append((idx, ans))

        total     = len(answers)
        num_width = len(str(total))
        ans_width = max(len(a) for _, a in answers)
        cell_width = num_width + 1 + ans_width + 2

        section    = doc.sections[0]
        usable_emu = section.page_width - section.left_margin - section.right_margin
        usable_pt  = usable_emu / 12700
        char_width_pt = 5.4
        max_chars  = int(usable_pt / char_width_pt)
        cols       = max(max_chars // cell_width, 1)

        for i in range(0, total, cols):
            chunk = answers[i:i + cols]
            parts = [f"{num}-{ans}".ljust(cell_width) for num, ans in chunk]
            line = "".join(parts).rstrip()

            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            run = p.add_run(line)
            run.font.size = Pt(9)
            run.font.name = "Courier New"
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.insert(0, rFonts)
            rFonts.set(qn("w:eastAsia"), "Courier New")

        # 图例说明
        if any(sq.answer_source == "ai"
               for q in questions for sq in q.sub_questions):
            p = doc.add_paragraph()
            run = p.add_run("* 标注表示该答案由 AI 补全，建议人工核对")
            run.font.size   = Pt(8)
            run.font.color.rgb = _AI_COLOR