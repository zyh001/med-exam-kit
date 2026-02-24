from __future__ import annotations
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from med_exam_toolkit.models import Question
from med_exam_toolkit.exporters import register
from med_exam_toolkit.exporters.base import BaseExporter

FONT_NAME = "宋体"
_AI_COLOR = RGBColor(0xCC, 0x77, 0x00)   # 橙色，与 exam/docx_exporter.py 保持一致


def _set_font(run, name: str = FONT_NAME, size: Pt | None = None):
    """同时设置中西文字体"""
    run.font.name = name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = size


@register("docx")
class DocxExporter(BaseExporter):

    def export(self, questions: list[Question], output_path: Path, **kwargs) -> None:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        fp = output_path.with_suffix(".docx")

        doc = Document()
        self._set_default_font(doc)

        doc.add_heading("医学考试题库", level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER

        for idx, q in enumerate(questions, 1):
            doc.add_heading(f"第{idx}题 [{q.mode}] {q.unit}", level=2)

            if q.stem:
                p   = doc.add_paragraph()
                run = p.add_run(f"【题干】{q.stem}")
                _set_font(run, size=Pt(10.5))

            if q.shared_options:
                p   = doc.add_paragraph()
                run = p.add_run("【共享选项】")
                run.bold = True
                _set_font(run)
                for opt in q.shared_options:
                    doc.add_paragraph(opt, style="List Bullet")

            for si, sq in enumerate(q.sub_questions, 1):
                prefix = f"({si}) " if len(q.sub_questions) > 1 else ""
                p   = doc.add_paragraph()
                run = p.add_run(f"{prefix}{sq.text}")
                run.bold = True
                _set_font(run)

                if sq.options and sq.options != q.shared_options:
                    for opt in sq.options:
                        doc.add_paragraph(opt, style="List Bullet")

                # 答案（AI 兜底时橙色 + 标注）
                ans = sq.eff_answer
                p   = doc.add_paragraph()
                run = p.add_run(f"答案: {ans}" if ans else "答案: —")
                if sq.answer_source == "ai":
                    run.font.color.rgb = _AI_COLOR
                    _set_font(run)
                    run2 = p.add_run("  (AI补全，建议核对)")
                    run2.font.size = Pt(8)
                    run2.font.color.rgb = _AI_COLOR
                    _set_font(run2)
                else:
                    run.font.color.rgb = RGBColor(0, 128, 0)
                    _set_font(run)

                if sq.rate:
                    run3 = p.add_run(f"  正确率: {sq.rate}")
                    _set_font(run3)

                # 解析（AI 兜底时同样橙色标注）
                dis     = sq.eff_discuss
                discuss = dis or q.discuss
                if discuss:
                    p   = doc.add_paragraph()
                    run = p.add_run(f"解析: {discuss}")
                    if sq.discuss_source == "ai":
                        run.font.color.rgb = _AI_COLOR
                        _set_font(run, size=Pt(9))
                        run2 = p.add_run("  (AI补全，建议核对)")
                        run2.font.size = Pt(8)
                        run2.font.color.rgb = _AI_COLOR
                        _set_font(run2)
                    else:
                        run.font.color.rgb = RGBColor(100, 100, 100)
                        _set_font(run, size=Pt(9))

            doc.add_paragraph("—" * 40)

        doc.save(fp)
        print(f"[INFO] DOCX 导出完成: {fp} ({len(questions)} 题)")

    @staticmethod
    def _set_default_font(doc: Document):
        """默认样式同时设置东亚字体"""
        style = doc.styles["Normal"]
        style.font.name = FONT_NAME
        style.font.size = Pt(10.5)
        rPr    = style.element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), FONT_NAME)